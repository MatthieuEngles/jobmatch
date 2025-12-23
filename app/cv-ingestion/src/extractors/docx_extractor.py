"""DOCX text extraction using python-docx."""

import logging
from io import BytesIO

from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text content from a DOCX file.

    Args:
        file_content: Raw bytes of the DOCX file.

    Returns:
        Extracted text as a single string with paragraphs separated.

    Raises:
        ValueError: If the DOCX cannot be read or is empty.
    """
    try:
        doc = Document(BytesIO(file_content))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        if not paragraphs:
            raise ValueError("No text could be extracted from DOCX")

        full_text = "\n".join(paragraphs)
        logger.info(f"Extracted {len(full_text)} chars from DOCX")

        return full_text

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}") from e
